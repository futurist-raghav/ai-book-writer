"""Book export generators (P2.7) - DOCX, PDF,EPUB support"""

from typing import List, Optional
from datetime import datetime
import re
from io import BytesIO


class MarkdownExporter:
    """Export book to Markdown format"""
    
    @staticmethod
    def export(book: dict, chapters: List[dict], include_metadata: bool = True) -> str:
        """
        Generate Markdown from book and chapters
        
        book: {title, description, author, status}
        chapters: [{title, content, chapter_number, part_name}]
        """
        lines = []
        
        if include_metadata:
            lines.append(f"# {book['title']}")
            lines.append("")
            if book.get('description'):
                lines.append(f"> {book['description']}")
                lines.append("")
            if book.get('authors'):
                author_str = ', '.join(book['authors'])
                lines.append(f"**Author:** {author_str}")
                lines.append("")
        
        current_part = None
        for chapter in chapters:
            part_name = chapter.get('part_name')
            
            # Add part heading if changed
            if part_name and part_name != current_part:
                current_part = part_name
                lines.append(f"## {part_name}")
                lines.append("")
            
            # Add chapter heading
            chapter_num = chapter.get('chapter_number', '')
            title = chapter.get('title', 'Untitled')
            if chapter_num:
                lines.append(f"### Chapter {chapter_num}: {title}")
            else:
                lines.append(f"### {title}")
            lines.append("")
            
            # Add chapter content
            content = chapter.get('content', '')
            if content:
                lines.append(content)
                lines.append("")
                lines.append("---")
                lines.append("")
        
        return "\n".join(lines)


class TextExporter:
    """Export book to plain text format"""
    
    @staticmethod
    def export(book: dict, chapters: List[dict], include_metadata: bool = True) -> str:
        """
        Generate plain text from book and chapters
        """
        lines = []
        
        if include_metadata:
            lines.append(book['title'].upper())
            lines.append("=" * len(book['title']))
            lines.append("")
            
            if book.get('description'):
                lines.append(book['description'])
                lines.append("")
            
            if book.get('authors'):
                author_str = ', '.join(book['authors'])
                lines.append(f"By {author_str}")
                lines.append("")
            
            lines.append("")
        
        current_part = None
        for chapter in chapters:
            part_name = chapter.get('part_name')
            
            if part_name and part_name != current_part:
                current_part = part_name
                lines.append("")
                lines.append(part_name.upper())
                lines.append("-" * len(part_name))
                lines.append("")
            
            chapter_num = chapter.get('chapter_number', '')
            title = chapter.get('title', 'Untitled')
            if chapter_num:
                lines.append(f"CHAPTER {chapter_num}")
                lines.append(title)
            else:
                lines.append(title.upper())
            
            lines.append("-" * 40)
            lines.append("")
            
            content = chapter.get('content', '')
            if content:
                # Clean up content - remove HTML tags if present
                content = re.sub(r'<[^>]+>', '', content)
                lines.append(content)
            
            lines.append("")
        
        return "\n".join(lines)


class DOCXExporter:
    """Export book to DOCX format (requires python-docx)"""
    
    @staticmethod
    def export(book: dict, chapters: List[dict], include_metadata: bool = True) -> bytes:
        """
        Generate DOCX from book and chapters
        
        Returns: bytes (DOCX file content)
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.style import WD_STYLE_TYPE
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            raise ImportError("python-docx is required for DOCX export")
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(book['title'], level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        if include_metadata:
            # Add metadata
            if book.get('description'):
                desc_para = doc.add_paragraph(book['description'])
                desc_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                desc_para.runs[0].italic = True
            
            if book.get('authors'):
                authors_para = doc.add_paragraph(f"By {', '.join(book['authors'])}")
                authors_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            doc.add_page_break()
        
        # Add chapters
        current_part = None
        for chapter in chapters:
            part_name = chapter.get('part_name')
            
            # Add part heading if needed
            if part_name and part_name != current_part:
                current_part = part_name
                doc.add_heading(part_name, level=1)
            
            # Add chapter heading
            chapter_num = chapter.get('chapter_number', '')
            title = chapter.get('title', 'Untitled')
            if chapter_num:
                doc.add_heading(f"Chapter {chapter_num}: {title}", level=2)
            else:
                doc.add_heading(title, level=2)
            
            # Add chapter content
            content = chapter.get('content', '')
            if content:
                # Remove HTML tags
                content = re.sub(r'<[^>]+>', '', content)
                
                # Split by newlines to preserve paragraph breaks
                for line in content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
            
            doc.add_page_break()
        
        # Save to bytes
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()


class EPUBExporter:
    """Export book to EPUB format (requires ebooklib)"""
    
    @staticmethod
    def export(book: dict, chapters: List[dict]) -> bytes:
        """
        Generate EPUB from book and chapters
        
        Requires: ebooklib
        """
        try:
            from ebooklib import epub
        except ImportError:
            raise ImportError("ebooklib is required for EPUB export")
        
        # Create EPUB book
        book_epub = epub.EpubBook()
        book_epub.set_identifier(f"book_{book['id']}_{datetime.utcnow().timestamp()}")
        book_epub.set_title(book['title'])
        
        if book.get('authors'):
            for author in book['authors']:
                book_epub.add_author(author)
        
        # Create chapters
        epub_chapters = []
        for idx, chapter in enumerate(chapters, 1):
            c = epub.EpubHtml(title=chapter.get('title', f'Chapter {idx}'),
                             file_name=f'chap_{idx:02d}.xhtml',
                             lang='en')
            
            # Create content
            content = chapter.get('content', '')
            content = re.sub(r'<[^>]+>', '', content)
            c.content = f'<h2>{chapter.get("title", f"Chapter {idx}")}</h2><p>{content.replace(chr(10), "</p><p>")}</p>'
            
            epub_chapters.append(c)
            book_epub.add_item(c)
        
        # Add navigation
        book_epub.toc = tuple(epub_chapters)
        book_epub.spine = ['nav'] + epub_chapters
        
        # Write to bytes
        output = BytesIO()
        epub.write_epub(output, book_epub, {})
        output.seek(0)
        return output.getvalue()
