"""
Export Background Tasks

Handles exporting books to various formats.
"""

import asyncio
import os
import uuid
from datetime import datetime, timezone
from io import BytesIO
from typing import Optional

from celery import shared_task
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from app.core.config import settings
from app.core.database import async_session_maker
from app.models.book import Book, BookChapter
from app.models.chapter import Chapter


async def _export_book_async(
    book_id: str,
    user_id: str,
    export_format: str,
    options: dict,
) -> dict:
    """
    Async implementation of book export.

    Args:
        book_id: UUID of the book.
        user_id: UUID of the user.
        export_format: Export format (pdf, epub, docx, html, markdown).
        options: Export options.

    Returns:
        Dictionary with export result info.
    """
    async with async_session_maker() as db:
        # Get book with chapters
        result = await db.execute(
            select(Book)
            .where(Book.id == book_id, Book.user_id == user_id)
            .options(
                selectinload(Book.chapter_associations).selectinload(
                    BookChapter.chapter
                )
            )
        )
        book = result.scalar_one_or_none()

        if not book:
            return {"error": "Book not found", "book_id": book_id}

        try:
            # Create export directory
            export_dir = os.path.join(settings.LOCAL_STORAGE_PATH, "exports", user_id)
            os.makedirs(export_dir, exist_ok=True)

            # Generate filename
            safe_title = "".join(
                c for c in book.title if c.isalnum() or c in " -_"
            ).strip()[:50]
            filename = f"{safe_title}_{uuid.uuid4().hex[:8]}.{export_format}"
            filepath = os.path.join(export_dir, filename)

            # Build content
            content = _build_book_content(book, options)

            # Export based on format
            if export_format == "markdown":
                _export_markdown(content, filepath)
            elif export_format == "html":
                _export_html(content, filepath, book)
            elif export_format == "pdf":
                _export_pdf(content, filepath, book, options)
            elif export_format == "docx":
                _export_docx(content, filepath, book, options)
            elif export_format == "epub":
                _export_epub(content, filepath, book, options)
            else:
                return {"error": f"Unsupported format: {export_format}"}

            # Get file size
            file_size = os.path.getsize(filepath)

            # Update book export info
            book.last_exported_at = datetime.now(timezone.utc)
            book.last_export_format = export_format
            await db.commit()

            return {
                "success": True,
                "book_id": book_id,
                "format": export_format,
                "filepath": filepath,
                "filename": filename,
                "file_size": file_size,
            }

        except Exception as e:
            return {
                "error": str(e),
                "book_id": book_id,
            }


def _build_book_content(book: Book, options: dict) -> dict:
    """Build structured book content for export."""
    content = {
        "title": book.title,
        "subtitle": book.subtitle,
        "author": book.author_name,
        "description": book.description,
        "front_matter": {},
        "chapters": [],
        "back_matter": {},
    }

    # Add front matter
    if options.get("include_front_matter", True):
        if book.dedication:
            content["front_matter"]["dedication"] = book.dedication
        if book.acknowledgments:
            content["front_matter"]["acknowledgments"] = book.acknowledgments
        if book.preface:
            content["front_matter"]["preface"] = book.preface
        if book.introduction:
            content["front_matter"]["introduction"] = book.introduction

    # Add chapters
    for assoc in sorted(book.chapter_associations, key=lambda x: x.order_index):
        chapter = assoc.chapter
        content["chapters"].append(
            {
                "number": chapter.chapter_number,
                "title": chapter.title,
                "subtitle": chapter.subtitle,
                "content": chapter.compiled_content or "",
                "part_number": assoc.part_number,
                "part_title": assoc.part_title,
            }
        )

    # Add back matter
    if options.get("include_back_matter", True):
        if book.epilogue:
            content["back_matter"]["epilogue"] = book.epilogue
        if book.afterword:
            content["back_matter"]["afterword"] = book.afterword
        if book.about_author:
            content["back_matter"]["about_author"] = book.about_author

    return content


def _export_markdown(content: dict, filepath: str) -> None:
    """Export book to Markdown format."""
    lines = []

    # Title page
    lines.append(f"# {content['title']}\n")
    if content.get("subtitle"):
        lines.append(f"## {content['subtitle']}\n")
    if content.get("author"):
        lines.append(f"*By {content['author']}*\n")
    lines.append("\n---\n")

    # Front matter
    for key, text in content.get("front_matter", {}).items():
        lines.append(f"\n## {key.replace('_', ' ').title()}\n")
        lines.append(f"{text}\n")

    # Table of contents
    lines.append("\n## Table of Contents\n")
    for ch in content.get("chapters", []):
        lines.append(f"- Chapter {ch['number']}: {ch['title']}\n")
    lines.append("\n---\n")

    # Chapters
    current_part = None
    for ch in content.get("chapters", []):
        if ch.get("part_number") and ch.get("part_number") != current_part:
            current_part = ch.get("part_number")
            lines.append(f"\n# Part {current_part}")
            if ch.get("part_title"):
                lines.append(f": {ch['part_title']}")
            lines.append("\n")

        lines.append(f"\n## Chapter {ch['number']}: {ch['title']}\n")
        if ch.get("subtitle"):
            lines.append(f"### {ch['subtitle']}\n")
        lines.append(f"\n{ch['content']}\n")

    # Back matter
    for key, text in content.get("back_matter", {}).items():
        lines.append(f"\n## {key.replace('_', ' ').title()}\n")
        lines.append(f"{text}\n")

    with open(filepath, "w", encoding="utf-8") as f:
        f.writelines(lines)


def _export_html(content: dict, filepath: str, book: Book) -> None:
    """Export book to HTML format."""
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{content['title']}</title>
    <style>
        body {{ font-family: Georgia, serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        h1 {{ text-align: center; }}
        h2 {{ margin-top: 2em; border-bottom: 1px solid #ccc; padding-bottom: 0.5em; }}
        .author {{ text-align: center; font-style: italic; }}
        .chapter {{ page-break-before: always; }}
        .front-matter, .back-matter {{ margin: 2em 0; }}
    </style>
</head>
<body>
    <h1>{content['title']}</h1>
"""
    if content.get("subtitle"):
        html += f"    <h2>{content['subtitle']}</h2>\n"
    if content.get("author"):
        html += f"    <p class='author'>By {content['author']}</p>\n"

    # Front matter
    for key, text in content.get("front_matter", {}).items():
        html += f"    <div class='front-matter'>\n"
        html += f"        <h2>{key.replace('_', ' ').title()}</h2>\n"
        html += f"        <p>{text.replace(chr(10), '</p><p>')}</p>\n"
        html += f"    </div>\n"

    # Chapters
    for ch in content.get("chapters", []):
        html += f"    <div class='chapter'>\n"
        html += f"        <h2>Chapter {ch['number']}: {ch['title']}</h2>\n"
        if ch.get("subtitle"):
            html += f"        <h3>{ch['subtitle']}</h3>\n"
        html += f"        <div>{ch['content'].replace(chr(10), '</p><p>')}</div>\n"
        html += f"    </div>\n"

    # Back matter
    for key, text in content.get("back_matter", {}).items():
        html += f"    <div class='back-matter'>\n"
        html += f"        <h2>{key.replace('_', ' ').title()}</h2>\n"
        html += f"        <p>{text.replace(chr(10), '</p><p>')}</p>\n"
        html += f"    </div>\n"

    html += "</body>\n</html>"

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html)


def _export_pdf(content: dict, filepath: str, book: Book, options: dict) -> None:
    """Export book to PDF format using ReportLab."""
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, PageBreak

    page_size = A4 if options.get("page_size") == "a4" else letter
    doc = SimpleDocTemplate(filepath, pagesize=page_size)
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "BookTitle",
        parent=styles["Title"],
        fontSize=24,
        spaceAfter=12,
    )
    chapter_style = ParagraphStyle(
        "ChapterTitle",
        parent=styles["Heading1"],
        fontSize=18,
        spaceBefore=24,
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "BookBody",
        parent=styles["Normal"],
        fontSize=options.get("font_size", 12),
        leading=options.get("font_size", 12) * 1.5,
    )

    story = []

    # Title page
    story.append(Paragraph(content["title"], title_style))
    if content.get("subtitle"):
        story.append(Paragraph(content["subtitle"], styles["Heading2"]))
    if content.get("author"):
        story.append(Spacer(1, 0.5 * inch))
        story.append(Paragraph(f"By {content['author']}", styles["Normal"]))
    story.append(PageBreak())

    # Front matter
    for key, text in content.get("front_matter", {}).items():
        story.append(Paragraph(key.replace("_", " ").title(), chapter_style))
        for para in text.split("\n\n"):
            if para.strip():
                story.append(Paragraph(para, body_style))
                story.append(Spacer(1, 0.2 * inch))
        story.append(PageBreak())

    # Chapters
    for ch in content.get("chapters", []):
        story.append(
            Paragraph(f"Chapter {ch['number']}: {ch['title']}", chapter_style)
        )
        if ch.get("subtitle"):
            story.append(Paragraph(ch["subtitle"], styles["Heading3"]))

        chapter_content = ch.get("content", "")
        for para in chapter_content.split("\n\n"):
            if para.strip():
                story.append(Paragraph(para, body_style))
                story.append(Spacer(1, 0.1 * inch))
        story.append(PageBreak())

    # Back matter
    for key, text in content.get("back_matter", {}).items():
        story.append(Paragraph(key.replace("_", " ").title(), chapter_style))
        for para in text.split("\n\n"):
            if para.strip():
                story.append(Paragraph(para, body_style))
                story.append(Spacer(1, 0.2 * inch))

    doc.build(story)


def _export_docx(content: dict, filepath: str, book: Book, options: dict) -> None:
    """Export book to DOCX format using python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()

    # Title
    title = doc.add_heading(content["title"], level=0)
    if content.get("subtitle"):
        doc.add_heading(content["subtitle"], level=1)
    if content.get("author"):
        doc.add_paragraph(f"By {content['author']}")

    doc.add_page_break()

    # Front matter
    for key, text in content.get("front_matter", {}).items():
        doc.add_heading(key.replace("_", " ").title(), level=1)
        doc.add_paragraph(text)
        doc.add_page_break()

    # Chapters
    for ch in content.get("chapters", []):
        doc.add_heading(f"Chapter {ch['number']}: {ch['title']}", level=1)
        if ch.get("subtitle"):
            doc.add_heading(ch["subtitle"], level=2)
        doc.add_paragraph(ch.get("content", ""))
        doc.add_page_break()

    # Back matter
    for key, text in content.get("back_matter", {}).items():
        doc.add_heading(key.replace("_", " ").title(), level=1)
        doc.add_paragraph(text)

    doc.save(filepath)


def _export_epub(content: dict, filepath: str, book: Book, options: dict) -> None:
    """Export book to EPUB format using ebooklib."""
    from ebooklib import epub

    ebook = epub.EpubBook()

    # Set metadata
    ebook.set_identifier(str(book.id))
    ebook.set_title(content["title"])
    ebook.set_language("en")
    if content.get("author"):
        ebook.add_author(content["author"])

    # Create chapters
    epub_chapters = []

    # Title page
    title_content = f"<h1>{content['title']}</h1>"
    if content.get("subtitle"):
        title_content += f"<h2>{content['subtitle']}</h2>"
    if content.get("author"):
        title_content += f"<p><em>By {content['author']}</em></p>"

    title_page = epub.EpubHtml(title="Title", file_name="title.xhtml")
    title_page.content = title_content
    ebook.add_item(title_page)
    epub_chapters.append(title_page)

    # Front matter
    for key, text in content.get("front_matter", {}).items():
        chapter = epub.EpubHtml(
            title=key.replace("_", " ").title(),
            file_name=f"{key}.xhtml",
        )
        chapter.content = f"<h1>{key.replace('_', ' ').title()}</h1><p>{text}</p>"
        ebook.add_item(chapter)
        epub_chapters.append(chapter)

    # Book chapters
    for ch in content.get("chapters", []):
        chapter = epub.EpubHtml(
            title=f"Chapter {ch['number']}: {ch['title']}",
            file_name=f"chapter_{ch['number']}.xhtml",
        )
        ch_content = f"<h1>Chapter {ch['number']}: {ch['title']}</h1>"
        if ch.get("subtitle"):
            ch_content += f"<h2>{ch['subtitle']}</h2>"
        ch_content += f"<div>{ch.get('content', '')}</div>"
        chapter.content = ch_content
        ebook.add_item(chapter)
        epub_chapters.append(chapter)

    # Back matter
    for key, text in content.get("back_matter", {}).items():
        chapter = epub.EpubHtml(
            title=key.replace("_", " ").title(),
            file_name=f"{key}.xhtml",
        )
        chapter.content = f"<h1>{key.replace('_', ' ').title()}</h1><p>{text}</p>"
        ebook.add_item(chapter)
        epub_chapters.append(chapter)

    # Add navigation
    ebook.toc = epub_chapters
    ebook.add_item(epub.EpubNcx())
    ebook.add_item(epub.EpubNav())
    ebook.spine = ["nav"] + epub_chapters

    # Write file
    epub.write_epub(filepath, ebook)


@shared_task(
    name="tasks.export_book",
    bind=True,
    autoretry_for=(Exception,),
    retry_backoff=True,
    retry_kwargs={"max_retries": 3},
)
def export_book(
    self,
    book_id: str,
    user_id: str,
    export_format: str,
    options: Optional[dict] = None,
) -> dict:
    """
    Celery task to export a book.

    Args:
        book_id: UUID of the book.
        user_id: UUID of the user.
        export_format: Export format.
        options: Export options.

    Returns:
        Dictionary with export result info.
    """
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        return loop.run_until_complete(
            _export_book_async(book_id, user_id, export_format, options or {})
        )
    finally:
        loop.close()
